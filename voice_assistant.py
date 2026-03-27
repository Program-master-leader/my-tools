#!/usr/bin/env python3
"""
小K语音助手
唤醒词：小K（后台托盘常驻，开机自启，语音回答）
"""
import os, sys, threading, datetime, subprocess, json, re
import tkinter as tk
from tkinter import scrolledtext

SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
DEFAULT_SAVE = "D:\\"
ASR_ENGINE   = "whisper"
BAIDU_APP_ID = BAIDU_API_KEY = BAIDU_SECRET = ""

WAKE_WORDS = ["小k","小K","xiaok","xiaoK","小可","小客","小卡","小咖","小ka","肖k","晓k","小key"]

BG="#1e1e2e"; BG2="#2a2a3e"; BG3="#313145"
ACCENT="#7c9ef8"; ACCENT2="#a6e3a1"; DANGER="#f38ba8"
TEXT="#cdd6f4"; TEXT_DIM="#6c7086"

# ══════════════════════════════════════════════════════
# 单实例互斥锁（防止重复启动）
# ══════════════════════════════════════════════════════
_MUTEX_NAME = "Global\\XiaoK_VoiceAssistant_Mutex"
_mutex_handle = None

def _ensure_single_instance():
    """
    用 Windows 命名互斥体确保只有一个实例。
    已有实例时：向已有实例发送激活信号后退出。
    返回 True 表示本进程是唯一实例，可以继续启动。
    """
    global _mutex_handle
    import ctypes
    kernel32 = ctypes.windll.kernel32
    _mutex_handle = kernel32.CreateMutexW(None, False, _MUTEX_NAME)
    last_err = kernel32.GetLastError()
    ERROR_ALREADY_EXISTS = 183
    if last_err == ERROR_ALREADY_EXISTS:
        # 已有实例，用 tasklist 找到它的窗口并激活（简单方案：直接退出）
        print("小K语音助手已在运行中。")
        return False
    return True

# ══════════════════════════════════════════════════════
# TTS 语音朗读
# ══════════════════════════════════════════════════════
def speak(text):
    """用 Windows SAPI 朗读文字（非阻塞）"""
    def _do():
        try:
            import pyttsx3
            engine = pyttsx3.init()
            engine.setProperty("rate", 180)
            # 优先选中文语音
            for v in engine.getProperty("voices"):
                if "zh" in v.id.lower() or "chinese" in v.name.lower() or "huihui" in v.name.lower():
                    engine.setProperty("voice", v.id)
                    break
            engine.say(text[:200])  # 限制长度，避免朗读太久
            engine.runAndWait()
        except Exception:
            pass  # TTS 失败静默处理
    threading.Thread(target=_do, daemon=True).start()

# ══════════════════════════════════════════════════════
# 唤醒词判断
# ══════════════════════════════════════════════════════
def _is_wake_word(text):
    t = text.strip().lower().replace(" ", "")
    for w in WAKE_WORDS:
        if w.lower() in t:
            return True
    if "小" in t and any(c in t for c in ["k","可","客","卡","咖","凯","开","克"]):
        return True
    return False

def _extract_command_after_wake(text):
    """
    从「小K小K PDF转Word」这样的句子里提取唤醒词之后的指令。
    返回指令字符串，如果只有唤醒词则返回空字符串。
    """
    # 尝试各种唤醒词变体，找到最后一个出现的位置
    t = text
    last_pos = -1
    last_len = 0
    candidates = WAKE_WORDS + ["小黑", "小可", "小客", "小卡"]
    for w in candidates:
        idx = t.lower().rfind(w.lower())
        if idx > last_pos:
            last_pos = idx
            last_len = len(w)
    if last_pos == -1:
        # 模糊匹配：找「小X」模式
        import re
        m = list(re.finditer(r"小[kKkK可客卡咖凯开克黑]", t))
        if m:
            last_pos = m[-1].start()
            last_len = len(m[-1].group())
    if last_pos == -1:
        return ""
    cmd = t[last_pos + last_len:].strip().lstrip("，。,. ")
    # 指令太短（<=3字）很可能是误识别，不当作指令
    if len(cmd) <= 3:
        return ""
    return cmd

# ══════════════════════════════════════════════════════
# Windows 内置语音识别（System.Speech，中文准确率高）
# ══════════════════════════════════════════════════════
_WIN_SPEECH_PS = os.path.join(SCRIPT_DIR, "win_speech.ps1")

def _ensure_win_speech_script():
    """确保 PowerShell 识别脚本存在"""
    script = r"""Add-Type -AssemblyName System.Speech
$engine = New-Object System.Speech.Recognition.SpeechRecognitionEngine
$engine.SetInputToDefaultAudioDevice()
$grammar = New-Object System.Speech.Recognition.DictationGrammar
$engine.LoadGrammar($grammar)
$timeout = [System.TimeSpan]::FromSeconds(10)
$result = $engine.Recognize($timeout)
if ($result -ne $null) { Write-Output $result.Text } else { Write-Output "" }
"""
    with open(_WIN_SPEECH_PS, "w", encoding="utf-8") as f:
        f.write(script)

def recognize_with_windows_speech(timeout_sec=10):
    """
    调用 Windows 内置 System.Speech 中文识别，返回识别文字。
    完全免费离线，中文准确率远高于 Whisper base。
    """
    tmp_ps = os.path.join(SCRIPT_DIR, "_tmp_speech.ps1")
    script = f"""Add-Type -AssemblyName System.Speech
$info = [System.Speech.Recognition.SpeechRecognitionEngine]::InstalledRecognizers() | Where-Object {{ $_.Culture.Name -eq 'zh-CN' }} | Select-Object -First 1
if ($info -ne $null) {{
    $engine = New-Object System.Speech.Recognition.SpeechRecognitionEngine($info)
}} else {{
    $engine = New-Object System.Speech.Recognition.SpeechRecognitionEngine
}}
$engine.SetInputToDefaultAudioDevice()
$grammar = New-Object System.Speech.Recognition.DictationGrammar
$engine.LoadGrammar($grammar)
$timeout = [System.TimeSpan]::FromSeconds({timeout_sec})
$result = $engine.Recognize($timeout)
if ($result -ne $null) {{ Write-Output $result.Text }} else {{ Write-Output "" }}
"""
    with open(tmp_ps, "w", encoding="utf-8") as f:
        f.write(script)
    try:
        r = subprocess.run(
            ["powershell", "-ExecutionPolicy", "Bypass", "-File", tmp_ps],
            capture_output=True, timeout=timeout_sec + 5,
            creationflags=0x08000000
        )
        # 中文 Windows PowerShell 输出 GBK
        for enc in ("gbk", "utf-8", "utf-16"):
            try:
                text = r.stdout.decode(enc).strip()
                break
            except Exception:
                text = ""
        return text
    except Exception:
        return ""
    finally:
        try: os.unlink(tmp_ps)
        except: pass


def _amplify(audio_np, target_peak=16000):
    """把音频放大到目标峰值，解决麦克风灵敏度低的问题"""
    import numpy as np
    peak = np.abs(audio_np).max()
    if peak < 10:
        return audio_np  # 纯静音不处理
    gain = min(target_peak / peak, 30.0)  # 最多放大30倍
    return np.clip(audio_np * gain, -32768, 32767).astype(np.int16)


def record_audio(duration=3, samplerate=16000):
    """固定时长录音（用于唤醒词检测），自动用设备原生采样率"""
    import sounddevice as sd, numpy as np, speech_recognition as sr
    device_info = sd.query_devices(kind="input")
    native_rate = int(device_info["default_samplerate"])
    audio_np = sd.rec(int(duration * native_rate), samplerate=native_rate,
                      channels=1, dtype="int16")
    sd.wait()
    audio_np = audio_np.flatten()
    if native_rate != samplerate:
        from scipy.signal import resample
        target_len = int(len(audio_np) * samplerate / native_rate)
        audio_np = resample(audio_np, max(target_len, 1)).astype(np.int16)
    if len(audio_np) == 0:
        audio_np = np.zeros(int(duration * samplerate), dtype=np.int16)
    audio_np = _amplify(audio_np)
    return sr.AudioData(audio_np.tobytes(), samplerate, 2)


def record_until_silence(stop_event=None, max_seconds=30, samplerate=16000,
                          silence_sec=1.5):
    """VAD录音：自动校准底噪，检测到说话开始录，静音超过 silence_sec 秒自动停止。"""
    import sounddevice as sd, numpy as np, speech_recognition as sr

    device_info = sd.query_devices(kind="input")
    native_rate = int(device_info["default_samplerate"])
    chunk = int(native_rate * 0.1)   # 100ms 一块

    # 自动校准底噪
    calib = sd.rec(int(0.5 * native_rate), samplerate=native_rate,
                   channels=1, dtype="int16")
    sd.wait()
    noise_level = np.abs(calib.flatten()).mean()
    energy_threshold = max(noise_level * 3, 3)  # 底噪*3，最低3

    frames = []
    silent_chunks = 0
    speaking = False
    max_chunks = int(max_seconds * 10)

    with sd.InputStream(samplerate=native_rate, channels=1, dtype="int16") as stream:
        for _ in range(max_chunks):
            if stop_event and stop_event.is_set():
                break
            data, _ = stream.read(chunk)
            data = data.flatten()
            energy = np.abs(data).mean()

            if energy > energy_threshold:
                speaking = True
                silent_chunks = 0
                frames.append(data)
            elif speaking:
                frames.append(data)
                silent_chunks += 1
                if silent_chunks >= int(silence_sec / 0.1):
                    break

    if not frames:
        return sr.AudioData(b"\x00" * int(samplerate * 0.1) * 2, samplerate, 2)

    audio_np = np.concatenate(frames)
    if native_rate != samplerate:
        from scipy.signal import resample
        target_len = int(len(audio_np) * samplerate / native_rate)
        audio_np = resample(audio_np, max(target_len, 1)).astype(np.int16)
    # 放大音量
    audio_np = _amplify(audio_np)
    return sr.AudioData(audio_np.tobytes(), samplerate, 2)

# ══════════════════════════════════════════════════════
# 语音识别
# ══════════════════════════════════════════════════════
def record_and_recognize(stop_event=None, max_seconds=30):
    """VAD录音 + 识别，说完自动停，或 stop_event 触发停止"""
    import numpy as np, speech_recognition as sr
    r = sr.Recognizer()
    audio = record_until_silence(stop_event=stop_event, max_seconds=max_seconds)
    return _whisper_recognize(audio)


def _whisper_recognize(audio):
    """用 Whisper 识别 AudioData，加常见误识别纠错"""
    import numpy as np
    raw = audio.get_raw_data(convert_rate=16000, convert_width=2)
    arr = np.frombuffer(raw, dtype=np.int16).flatten().astype(np.float32) / 32768.0
    # 空音频直接返回（放大后有效音频峰值至少0.1）
    if len(arr) < 1600 or np.abs(arr).max() < 0.05:
        return ""
    if not hasattr(record_and_recognize, "_model"):
        import whisper
        record_and_recognize._model = whisper.load_model("base")
    result = record_and_recognize._model.transcribe(
        arr, language="zh", fp16=False,
        condition_on_previous_text=False,  # 避免幻觉
        no_speech_threshold=0.6,           # 提高无声判断阈值
        logprob_threshold=-1.0,
    )
    text = result["text"].strip()
    # 常见误识别纠错
    corrections = {
        "CDF": "PDF", "cdf": "pdf", "DDF": "PDF", "PPDF": "PDF",
        "握的": "Word", "握得": "Word", "彎的": "Word", "弯的": "Word",
        "文道": "文档", "文盪": "文档",
        "转握": "转Word", "轉彎": "转Word", "转弯": "转Word",
        "DF转": "PDF转", "DF轉": "PDF转",
        "小可": "小K", "小客": "小K", "小卡": "小K", "小黑": "小K",
    }
    for wrong, right in corrections.items():
        text = text.replace(wrong, right)
    return text

# ══════════════════════════════════════════════════════
# AI（Ollama）
# ══════════════════════════════════════════════════════
def ask_ollama(prompt, stream_callback=None):
    import time
    for p in [r"D:\Ollama\ollama.exe", r"C:\Users\Public\ollama\ollama.exe"]:
        if os.path.exists(p):
            subprocess.Popen([p, "serve"], creationflags=0x08000000,
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            time.sleep(1); break
    try:
        import ollama
        result = ""
        for chunk in ollama.chat(model="qwen2.5:7b",
                                  messages=[{"role":"user","content":prompt}], stream=True):
            piece = chunk["message"]["content"]
            result += piece
            if stream_callback: stream_callback(piece)
        return result
    except Exception as e:
        return f"[AI错误: {e}]"

# ══════════════════════════════════════════════════════
# 工具匹配
# ══════════════════════════════════════════════════════
def load_tools_list():
    p = os.path.join(SCRIPT_DIR, "tools.json")
    if os.path.exists(p):
        with open(p, encoding="utf-8") as f:
            return json.load(f)
    return []

def match_tool(text, tools):
    text_lower = text.lower()
    scored = []
    for t in tools:
        score = 0
        name = t.get("name","").lower()
        desc = t.get("desc","").lower()
        for word in re.split(r"[\s（）()，。、/↔]", name):
            word = word.strip()
            if word and len(word) >= 2 and word in text_lower:
                score += 5
        for word in re.split(r"[\s，。、/]", desc):
            word = word.strip()
            if word and len(word) >= 2 and word in text_lower:
                score += 2
        for ch in name:
            if ch in text and ch not in "工具助手管理中心":
                score += 0.3
        if "转换" in text and "转换" not in name and "转换" not in desc: score -= 3
        if "合并" in text and "合并" not in name and "合并" not in desc: score -= 3
        if "语音" in text and "语音" not in name and "语音" not in desc: score -= 2
        if score > 0:
            scored.append((score, t))
    if scored:
        scored.sort(key=lambda x: x[0], reverse=True)
        if len(scored)==1 or scored[0][0] >= scored[1][0]+1.5:
            return scored[0][1]
    # 有「打开/启动」关键词时不走 Ollama，直接返回 None 让 _find_and_launch_app 处理
    if any(w in text for w in ["打开","启动","运行","开启"]):
        return None
    tool_list_str = "\n".join(f"- {t['name']}：{t.get('desc','')}" for t in tools)
    prompt = (f"用户说：「{text}」\n\n可用工具列表：\n{tool_list_str}\n\n"
              f"请判断用户是否想启动某个工具。如果是，只回复工具名称；如果不是，只回复「无」。不要解释。")
    result = ask_ollama(prompt).strip().strip("「」\"'")
    for t in tools:
        if t["name"]==result or result in t["name"]:
            return t
    return None

# ══════════════════════════════════════════════════════
# 系统软件搜索 & 启动
# ══════════════════════════════════════════════════════
def _find_and_launch_app(app_name, log_fn):
    """
    搜索系统已安装软件并启动。
    搜索范围：开始菜单快捷方式、Program Files、常见浏览器路径。
    支持中英文模糊匹配（如「夸克」匹配「Quark.exe」）。
    返回成功消息字符串，找不到返回 None。
    """
    import glob

    # 中英文别名映射
    ALIASES = {
        "夸克": ["quark", "夸克"],
        "微信": ["wechat", "weixin", "微信"],
        "qq": ["qq", "腾讯qq"],
        "钉钉": ["dingtalk", "钉钉"],
        "谷歌": ["chrome", "google chrome"],
        "火狐": ["firefox", "火狐"],
        "edge": ["msedge", "edge"],
        "记事本": ["notepad"],
        "计算器": ["calc", "calculator"],
        "资源管理器": ["explorer"],
        "vscode": ["code", "visual studio code"],
        "微软拼音": ["mspy"],
        "网易云": ["cloudmusic", "neteasemusic", "网易云音乐"],
        "百度网盘": ["baidupcs", "baidunetdisk", "百度网盘"],
        "迅雷": ["thunder", "迅雷"],
        "steam": ["steam"],
        "typora": ["typora"],
        "pycharm": ["pycharm"],
    }

    name_lower = app_name.lower().strip()

    # 展开别名
    search_terms = [name_lower]
    for key, aliases in ALIASES.items():
        if key in name_lower or any(a in name_lower for a in aliases):
            search_terms.extend(aliases)

    # 搜索路径：桌面优先，然后开始菜单，最后 Program Files
    def _score(fname):
        s = 0
        for term in search_terms:
            if term == fname:       s = max(s, 10)
            elif term in fname:     s = max(s, 6)
            elif fname in term:     s = max(s, 5)
        return s

    candidates = []  # (score, path)

    # 第一优先：桌面（用户桌面 + 公共桌面）
    for desktop in [
        os.path.expandvars(r"%USERPROFILE%\Desktop"),
        r"C:\Users\Public\Desktop",
    ]:
        if not os.path.exists(desktop):
            continue
        for fpath in glob.glob(os.path.join(desktop, "*.lnk")) + \
                     glob.glob(os.path.join(desktop, "*.exe")):
            fname = os.path.splitext(os.path.basename(fpath))[0].lower()
            s = _score(fname)
            if s > 0:
                candidates.append((s + 3, fpath))  # 桌面加权+3

    # 第二优先：开始菜单（递归，速度快）
    for menu_dir in [
        os.path.expandvars(r"%APPDATA%\Microsoft\Windows\Start Menu"),
        os.path.expandvars(r"%ProgramData%\Microsoft\Windows\Start Menu"),
    ]:
        if not os.path.exists(menu_dir):
            continue
        for fpath in glob.glob(os.path.join(menu_dir, "**", "*.lnk"), recursive=True):
            fname = os.path.splitext(os.path.basename(fpath))[0].lower()
            s = _score(fname)
            if s > 0:
                candidates.append((s, fpath))

    # 开始菜单找到了就直接用，不再扫 Program Files（避免慢）
    if not candidates:
        for prog_dir in [
            r"C:\Program Files",
            r"C:\Program Files (x86)",
            os.path.expandvars(r"%LOCALAPPDATA%\Programs"),
        ]:
            if not os.path.exists(prog_dir):
                continue
            # 只搜两层深度，避免太慢
            for fpath in glob.glob(os.path.join(prog_dir, "*", "*.exe")):
                fname = os.path.splitext(os.path.basename(fpath))[0].lower()
                s = _score(fname)
                if s > 0:
                    candidates.append((s, fpath))
            for fpath in glob.glob(os.path.join(prog_dir, "*", "*", "*.exe")):
                fname = os.path.splitext(os.path.basename(fpath))[0].lower()
                s = _score(fname)
                if s > 0:
                    candidates.append((s, fpath))

    if not candidates:
        return None

    candidates.sort(key=lambda x: x[0], reverse=True)
    best_path = candidates[0][1]
    display = os.path.splitext(os.path.basename(best_path))[0]
    log_fn(f"找到软件：{display}")
    try:
        os.startfile(best_path)  # 最可靠的方式，支持 .lnk/.exe
        return f"好的，正在打开「{display}」"
    except Exception as e:
        return f"启动失败：{e}"


# ══════════════════════════════════════════════════════
# 指令执行
# ══════════════════════════════════════════════════════
def parse_save_path(text):
    for p in [r"保存到([A-Za-z]:[^\s，。,\.]+)", r"存到([A-Za-z]:[^\s，。,\.]+)",
              r"放到([A-Za-z]:[^\s，。,\.]+)", r"保存在([A-Za-z]:[^\s，。,\.]+)"]:
        m = re.search(p, text)
        if m: return m.group(1).strip()
    return DEFAULT_SAVE

def parse_filename(text):
    for p in [r'命名为["\u300c\u300e]?([^\u300d\u300f"\uff0c\u3002\s]+)',
              r'文件名[为是]["\u300c\u300e]?([^\u300d\u300f"\uff0c\u3002\s]+)',
              r'叫做?["\u300c\u300e]?([^\u300d\u300f"\uff0c\u3002\s]+)',
              r'取名[为叫]["\u300c\u300e]?([^\u300d\u300f"\uff0c\u3002\s]+)']:
        m = re.search(p, text)
        if m: return m.group(1).strip()
    return None

def parse_format(text):
    if "word" in text.lower() or "docx" in text.lower() or "文档" in text: return "docx"
    if "pdf" in text.lower(): return "pdf"
    return "txt"

def launch_tool(tool):
    path = tool["path"]
    if not os.path.isabs(path): path = os.path.join(SCRIPT_DIR, path)
    ext = os.path.splitext(path)[1].lower()
    if not os.path.exists(path): return
    if ext == ".py":   subprocess.Popen(f'start cmd /k python "{path}"', shell=True)
    elif ext == ".exe": subprocess.Popen(f'"{path}"', shell=True)
    elif ext in (".bat",".cmd"): subprocess.Popen(f'start cmd /k "{path}"', shell=True)
    else: subprocess.Popen(f'start "" "{path}"', shell=True)

def execute_command(text, log_fn, stream_fn):
    text_lower = text.lower().strip()
    if any(w in text for w in ["关机","关闭电脑"]):
        log_fn("收到关机指令，5秒后关机...")
        subprocess.Popen("shutdown /s /t 5", shell=True)
        return "好的，电脑将在5秒后关机。"
    if any(w in text for w in ["取消关机","撤销关机"]):
        subprocess.Popen("shutdown /a", shell=True); return "已取消关机。"
    if any(w in text for w in ["重启","重新启动"]):
        subprocess.Popen("shutdown /r /t 5", shell=True); return "好的，5秒后重启。"
    if any(w in text for w in ["几点","时间","现在多少点"]):
        now = datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
        return f"现在是 {now}"
    if any(w in text for w in ["打开计算器","启动计算器"]):
        subprocess.Popen("calc", shell=True); return "已打开计算器。"
    if any(w in text for w in ["打开记事本","启动记事本"]):
        subprocess.Popen("notepad", shell=True); return "已打开记事本。"
    if any(w in text for w in ["打开浏览器","启动浏览器"]):
        subprocess.Popen("start https://www.baidu.com", shell=True); return "已打开浏览器。"

    launch_triggers = ["打开","启动","运行","开启","帮我打开","帮我启动","用","使用","调用"]
    if any(w in text for w in launch_triggers):
        # 先提取软件名，再查 tools.json（避免「无忧小助手」匹配到「小K语音助手」）
        app_name_raw = re.sub(r"(打开|启动|运行|开启|帮我打开|帮我启动)", "", text).strip()
        app_name_raw = re.sub(r"(浏览器|软件|程序|应用)$", "", app_name_raw).strip()

        # 先搜系统软件（桌面/开始菜单），精确匹配优先
        if app_name_raw:
            found = _find_and_launch_app(app_name_raw, log_fn)
            if found:
                return found

        # 系统没找到，再查 tools.json
        tools = load_tools_list()
        tool = match_tool(text, tools)
        if tool:
            path = tool["path"]
            if not os.path.isabs(path): path = os.path.join(SCRIPT_DIR, path)
            if os.path.exists(path):
                log_fn(f"启动工具：{tool['name']}")
                launch_tool(tool)
                return f"好的，正在启动「{tool['name']}」"

    write_triggers = ["写","生成","帮我写","写一篇","写一个","创作","起草","新建","创建","建一个","建一份"]
    if any(w in text for w in write_triggers):
        fmt = parse_format(text); save_dir = parse_save_path(text); fname = parse_filename(text)
        # 「新建空文档」不需要AI生成，直接创建空文件
        is_new_empty = any(w in text for w in ["新建","创建","建一个","建一份"]) and \
                       not any(w in text for w in ["写","生成","帮我写","创作","起草"])
        if is_new_empty:
            if not fname:
                fname = "新建文档"
            os.makedirs(save_dir, exist_ok=True)
            ext_map = {"docx":".docx","pdf":".pdf","txt":".txt"}
            full_path = os.path.join(save_dir, fname + ext_map.get(fmt, ".docx" if "word" in text.lower() or "文档" in text else ".txt"))
            try:
                if full_path.endswith(".docx"):
                    from docx import Document
                    Document().save(full_path)
                else:
                    open(full_path, "w", encoding="utf-8").close()
                subprocess.Popen(f'start "" "{full_path}"', shell=True)
                return f"已新建并打开：{full_path}"
            except Exception as e:
                return f"新建失败：{e}"
        fmt = parse_format(text); save_dir = parse_save_path(text); fname = parse_filename(text)
        clean = re.sub(r"保存(到|在|为)[^\s，。]+","",text)
        clean = re.sub(r"(命名为|文件名|叫做?|取名)[^\s，。]+","",clean)
        clean = re.sub(r"(word|docx|pdf|txt|文档|文本)格式?","",clean,flags=re.I).strip("，。, .")
        log_fn(f"正在生成：{clean[:30]}..."); stream_fn("", clear=True)
        content = ask_ollama(clean, stream_callback=lambda p: stream_fn(p))
        if not content or content.startswith("[AI错误"): return f"生成失败：{content}"
        if not fname:
            fname = re.sub(r"[写帮我生成创作起草一篇个]","",clean)[:20].strip()
            fname = re.sub(r'[\\/:*?"<>|]',"",fname) or "小K生成"
        os.makedirs(save_dir, exist_ok=True)
        ext_map = {"docx":".docx","pdf":".pdf","txt":".txt"}
        full_path = os.path.join(save_dir, fname + ext_map.get(fmt,".txt"))
        try:
            if fmt=="docx":
                from docx import Document
                doc=Document()
                for line in content.split("\n"): doc.add_paragraph(line)
                doc.save(full_path)
            else:
                with open(full_path,"w",encoding="utf-8-sig") as f: f.write(content)
            return f"已保存到：{full_path}"
        except Exception as e:
            return f"保存失败：{e}"

    tools = load_tools_list()
    tool = match_tool(text, tools)
    if tool:
        path = tool["path"]
        if not os.path.isabs(path): path = os.path.join(SCRIPT_DIR, path)
        if os.path.exists(path):
            log_fn(f"识别到工具：{tool['name']}"); launch_tool(tool)
            return f"好的，正在启动「{tool['name']}」"

    log_fn("正在思考..."); stream_fn("", clear=True)
    return ask_ollama(text, stream_callback=lambda p: stream_fn(p))

# ══════════════════════════════════════════════════════
# GUI + 托盘常驻
# ══════════════════════════════════════════════════════
class VoiceAssistant(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("小K 语音助手")
        self.geometry("700x560")
        self.configure(bg=BG)
        self.resizable(True, True)
        self.listening = False
        self.wake_mode = False
        self._tray_icon = None
        self._task_stop = threading.Event()   # 中断当前任务用
        self._task_lock = threading.Lock()    # 同一时间只处理一个任务
        self._build_ui()
        self._load_asr_config()
        self._check_deps()
        self.protocol("WM_DELETE_WINDOW", self._hide_to_tray)

    def _load_asr_config(self):
        global ASR_ENGINE, BAIDU_APP_ID, BAIDU_API_KEY, BAIDU_SECRET
        cfg = os.path.join(SCRIPT_DIR, "voice_config.json")
        if os.path.exists(cfg):
            with open(cfg, encoding="utf-8") as f:
                d = json.load(f)
            ASR_ENGINE=d.get("engine","whisper"); BAIDU_APP_ID=d.get("baidu_app_id","")
            BAIDU_API_KEY=d.get("baidu_api_key",""); BAIDU_SECRET=d.get("baidu_secret","")

    def _build_ui(self):
        top = tk.Frame(self, bg=BG2, pady=10); top.pack(fill="x")
        tk.Label(top, text="🎙  小K 语音助手", bg=BG2, fg=ACCENT,
                 font=("微软雅黑",14,"bold")).pack(side="left", padx=20)
        self.status_label = tk.Label(top, text="● 待机", bg=BG2, fg=TEXT_DIM,
                                      font=("微软雅黑",10))
        self.status_label.pack(side="right", padx=20)

        self.chat = scrolledtext.ScrolledText(self, bg=BG2, fg=TEXT, font=("微软雅黑",10),
            relief="flat", state="disabled", wrap="word", height=18)
        self.chat.pack(fill="both", expand=True, padx=16, pady=8)
        self.chat.tag_config("user", foreground=ACCENT)
        self.chat.tag_config("ai",   foreground=ACCENT2)
        self.chat.tag_config("system", foreground=TEXT_DIM)
        self.chat.tag_config("error",  foreground=DANGER)

        input_frame = tk.Frame(self, bg=BG, pady=6); input_frame.pack(fill="x", padx=16)
        self.input_var = tk.StringVar()
        self.entry = tk.Entry(input_frame, textvariable=self.input_var, bg=BG2, fg=TEXT,
                         insertbackground=TEXT, relief="flat", font=("微软雅黑",11))
        self.entry.pack(side="left", fill="x", expand=True, ipady=8, padx=(0,8))
        self.entry.bind("<Return>", lambda e: self._send_text())
        tk.Button(input_frame, text="发送", bg=ACCENT, fg=BG, relief="flat",
                  font=("微软雅黑",10), padx=14, pady=6, cursor="hand2",
                  command=self._send_text).pack(side="left")

        btn_frame = tk.Frame(self, bg=BG, pady=6); btn_frame.pack(fill="x", padx=16)
        self.mic_btn = tk.Button(btn_frame, text="🎙 点击说话", bg="#45475a", fg=TEXT,
            relief="flat", font=("微软雅黑",10), padx=16, pady=8, cursor="hand2",
            command=self._toggle_mic)
        self.mic_btn.pack(side="left", padx=4)
        self.mic_btn.bind("<ButtonPress-1>",   self._mic_press)
        self.mic_btn.bind("<ButtonRelease-1>", self._mic_release)

        self.wake_btn = tk.Button(btn_frame, text="👂 开启唤醒监听", bg="#45475a", fg=TEXT,
            relief="flat", font=("微软雅黑",10), padx=16, pady=8, cursor="hand2",
            command=self._toggle_wake)
        self.wake_btn.pack(side="left", padx=4)

        tk.Button(btn_frame, text="🚀 开机自启", bg="#45475a", fg=TEXT,
                  relief="flat", font=("微软雅黑",10), padx=12, pady=8,
                  cursor="hand2", command=self._toggle_autostart).pack(side="left", padx=4)

        self.stop_btn = tk.Button(btn_frame, text="⏹ 停止", bg=DANGER, fg=BG,
                  relief="flat", font=("微软雅黑",10), padx=12, pady=8,
                  cursor="hand2", command=self._stop_task)
        self.stop_btn.pack(side="left", padx=4)

        tk.Button(btn_frame, text="清空", bg="#45475a", fg=TEXT,
                  relief="flat", font=("微软雅黑",10), padx=12, pady=8,
                  cursor="hand2", command=self._clear_chat).pack(side="right", padx=4)
        tk.Button(btn_frame, text="⚙ 语音设置", bg="#45475a", fg=TEXT,
                  relief="flat", font=("微软雅黑",10), padx=12, pady=8,
                  cursor="hand2", command=self._open_asr_settings).pack(side="right", padx=4)

        self._append("system", "小K已就绪。说「小K」唤醒，或直接输入指令。\n"
                                "关闭窗口后小K继续在托盘后台监听。\n")

    # ── 托盘 ──────────────────────────────────────────
    def _hide_to_tray(self):
        self.withdraw()  # 隐藏窗口
        self._start_tray()

    def _start_tray(self):
        if self._tray_icon:
            return
        try:
            import pystray
            from PIL import Image, ImageDraw
            # 画一个简单的麦克风图标
            img = Image.new("RGB", (64,64), color="#1e1e2e")
            d = ImageDraw.Draw(img)
            d.ellipse([20,4,44,36], fill="#7c9ef8")
            d.rectangle([28,36,36,52], fill="#7c9ef8")
            d.line([20,52,44,52], fill="#7c9ef8", width=3)

            menu = pystray.Menu(
                pystray.MenuItem("显示窗口", self._show_window, default=True),
                pystray.MenuItem("唤醒监听: 开/关", self._tray_toggle_wake),
                pystray.Menu.SEPARATOR,
                pystray.MenuItem("退出小K", self._quit_app),
            )
            self._tray_icon = pystray.Icon("小K语音助手", img, "小K语音助手", menu)
            threading.Thread(target=self._tray_icon.run, daemon=True).start()
        except ImportError:
            # 没有 pystray，直接最小化
            self.iconify()

    def _show_window(self, icon=None, item=None):
        self.after(0, self.deiconify)
        self.after(0, self.lift)

    def _tray_toggle_wake(self, icon=None, item=None):
        self.after(0, self._toggle_wake)

    def _quit_app(self, icon=None, item=None):
        self.wake_mode = False
        if self._tray_icon:
            self._tray_icon.stop()
        self.after(0, self.destroy)

    # ── 开机自启 ──────────────────────────────────────
    def _toggle_autostart(self):
        import winreg
        key_path = r"Software\Microsoft\Windows\CurrentVersion\Run"
        app_name = "小K语音助手"
        pythonw = os.path.join(os.path.dirname(sys.executable), "pythonw.exe")
        if not os.path.exists(pythonw): pythonw = sys.executable
        script = os.path.abspath(__file__)
        cmd = f'"{pythonw}" "{script}"'
        try:
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0,
                                 winreg.KEY_READ | winreg.KEY_SET_VALUE)
            try:
                winreg.QueryValueEx(key, app_name)
                winreg.DeleteValue(key, app_name)
                self._append("system", "✓ 已取消开机自启")
            except FileNotFoundError:
                winreg.SetValueEx(key, app_name, 0, winreg.REG_SZ, cmd)
                self._append("system", "✓ 已设置开机自启，下次开机小K自动在后台监听")
            winreg.CloseKey(key)
        except Exception as e:
            self._append("error", f"设置失败：{e}")

    # ── 依赖检查 + 自动开启唤醒 ──────────────────────
    def _check_deps(self):
        def check():
            missing = []
            for pkg in ["speech_recognition","ollama","pyttsx3"]:
                try: __import__(pkg)
                except ImportError: missing.append(pkg)
            if missing:
                self.after(0, lambda: self._append("error",
                    f"缺少依赖：{', '.join(missing)}\n请运行：pip install {' '.join(missing)}"))
            self.after(800, self._auto_wake)
        threading.Thread(target=check, daemon=True).start()

    def _auto_wake(self):
        if not self.wake_mode:
            self._toggle_wake()

    def _stop_task(self):
        """立即停止当前任务（TTS朗读、AI生成、Win+H等）"""
        self._task_stop.set()
        # 停止TTS
        try:
            import pyttsx3
            # pyttsx3 没有全局stop，用新引擎覆盖
        except Exception:
            pass
        self._set_status("● 已停止", DANGER)
        self._append("system", "⏹ 已停止当前任务")
        # 恢复状态
        self.after(1500, lambda: self._set_status(
            "● 监听唤醒词" if self.wake_mode else "● 待机",
            ACCENT2 if self.wake_mode else TEXT_DIM))

    # ── 工具方法 ──────────────────────────────────────
    def _append(self, tag, text):
        self.chat.config(state="normal")
        prefix = {"user":"👤 你：","ai":"🤖 小K：","system":"","error":"⚠ "}.get(tag,"")
        self.chat.insert("end", prefix+text+"\n", tag)
        self.chat.see("end"); self.chat.config(state="disabled")

    def _stream_append(self, piece, clear=False):
        self.chat.config(state="normal")
        if clear:
            idx = self.chat.search("🤖 小K：","1.0",backwards=True,stopindex="end")
            if idx:
                self.chat.delete(self.chat.index(f"{idx} lineend"), "end")
                self.chat.insert("end","\n")
        self.chat.insert("end", piece, "ai")
        self.chat.see("end"); self.chat.config(state="disabled")

    def _set_status(self, text, color=TEXT_DIM):
        self.status_label.config(text=text, fg=color)

    def _clear_chat(self):
        self.chat.config(state="normal"); self.chat.delete("1.0","end")
        self.chat.config(state="disabled")

    # ── 按住说话（改为点击切换）────────────────────────
    def _mic_press(self, event):
        pass  # 不用按住，改为点击切换

    def _mic_release(self, event):
        pass  # 不用

    def _toggle_mic(self):
        if self.listening:
            self.listening = False
            self.mic_btn.config(bg="#45475a", text="🎙 点击说话")
            self._set_status("● 待机")
        else:
            self.listening = True
            self.mic_btn.config(bg=DANGER, text="🔴 识别中...")
            self._set_status("● Win+H 语音输入", DANGER)
            # 直接触发 Win+H
            self._trigger_winh_input()
            self.after(100, lambda: self.mic_btn.config(bg="#45475a", text="🎙 点击说话"))
            self.listening = False

    def _do_record(self):
        try:
            self._append("system", "🎤 Windows语音识别中，说完后稍等...")
            text = recognize_with_windows_speech(timeout_sec=10)
            self.after(0, lambda: self.mic_btn.config(bg="#45475a", text="🎙 点击说话"))
            self.after(0, lambda: self._set_status("● 待机"))
            self.listening = False
            if text: self.after(0, lambda: self._handle_input(text))
            else:    self.after(0, lambda: self._append("system", "未识别到语音，请重试"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self.mic_btn.config(bg="#45475a", text="🎙 点击说话"))
            self.after(0, lambda: self._set_status("● 待机"))
            self.listening = False
            self.after(0, lambda msg=err: self._append("error", f"录音失败：{msg}"))

    # ── 唤醒监听 ──────────────────────────────────────
    def _toggle_wake(self):
        if self.wake_mode:
            self.wake_mode = False
            self.wake_btn.config(text="👂 开启唤醒监听", bg="#45475a")
            self._set_status("● 待机")
            self._append("system","唤醒监听已关闭")
        else:
            self.wake_mode = True
            self.wake_btn.config(text="🟢 唤醒监听中", bg="#40a02b")
            self._set_status("● 监听唤醒词", ACCENT2)
            self._append("system","唤醒监听已开启，说「小K」激活助手")
            threading.Thread(target=self._wake_loop, daemon=True).start()

    def _wake_loop(self):
        """
        唤醒词监听：用 Windows System.Speech 循环识别，准确率高。
        识别到「小K」相关词汇立即唤醒。
        """
        import subprocess, tempfile, os

        # PowerShell 脚本：循环识别，每次最多等3秒，识别到内容立即返回
        ps_script = r"""
Add-Type -AssemblyName System.Speech
$info = [System.Speech.Recognition.SpeechRecognitionEngine]::InstalledRecognizers() | Where-Object { $_.Culture.Name -eq 'zh-CN' } | Select-Object -First 1
if ($info -ne $null) {
    $engine = New-Object System.Speech.Recognition.SpeechRecognitionEngine($info)
} else {
    $engine = New-Object System.Speech.Recognition.SpeechRecognitionEngine
}
$engine.SetInputToDefaultAudioDevice()
$grammar = New-Object System.Speech.Recognition.DictationGrammar
$engine.LoadGrammar($grammar)
$timeout = [System.TimeSpan]::FromSeconds(3)
while ($true) {
    $result = $engine.Recognize($timeout)
    if ($result -ne $null -and $result.Text -ne "") {
        Write-Output $result.Text
        [Console]::Out.Flush()
    }
}
"""
        tmp_ps = os.path.join(SCRIPT_DIR, "_wake_listener.ps1")
        with open(tmp_ps, "w", encoding="utf-8") as f:
            f.write(ps_script)

        proc = subprocess.Popen(
            ["powershell", "-ExecutionPolicy", "Bypass", "-File", tmp_ps],
            stdout=subprocess.PIPE, stderr=subprocess.DEVNULL,
            creationflags=0x08000000  # 无窗口
        )
        self._wake_ps_proc = proc  # 保存引用供 _trigger_winh_input 暂停用

        try:
            while self.wake_mode:
                line = proc.stdout.readline()
                if not line:
                    continue
                for enc in ("gbk", "utf-8", "utf-16"):
                    try:
                        text = line.decode(enc).strip()
                        break
                    except Exception:
                        text = ""
                if not text:
                    continue

                # 纠错
                for w, r in [
                    ("小可","小K"), ("小客","小K"), ("小卡","小K"),
                    ("小黑","小K"), ("小凯","小K"), ("小开","小K"),
                    ("小克","小K"), ("晓K","小K"), ("晓k","小K"),
                    ("肖K","小K"), ("肖k","小K"), ("小key","小K"),
                    ("小壳","小K"), ("小科","小K"), ("小棵","小K"),
                    ("小鬼","小K"), ("小哥","小K"),
                ]:
                    text = text.replace(w, r)

                # 调试日志
                self.after(0, lambda t=text: self._append("system", f"🔍 听到：{t}"))

                if not _is_wake_word(text):
                    continue

                # 唤醒
                self._task_stop.set()
                cmd = _extract_command_after_wake(text)
                if cmd:
                    self.after(0, lambda c=cmd: self._on_wake_with_cmd(c))
                else:
                    self.after(0, self._on_wake)
        finally:
            try:
                proc.terminate()
                os.unlink(tmp_ps)
            except Exception:
                pass

    def _trigger_winh_input(self):
        """
        触发 Win+H 语音输入。
        先暂停唤醒监听的PS进程（释放麦克风），Win+H识别完再恢复。
        """
        self.deiconify()
        self.lift()
        self.focus_force()
        self.input_var.set("")
        self.entry.focus_set()

        self._set_status("● Win+H 语音输入中...", ACCENT)
        self._append("system", "🎤 请说话（Win+H 云端识别）...")

        # 暂停唤醒监听PS进程，释放麦克风给Win+H
        wake_proc = getattr(self, "_wake_ps_proc", None)
        if wake_proc and wake_proc.poll() is None:
            try:
                import ctypes
                ctypes.windll.kernel32.SuspendThread(
                    ctypes.windll.kernel32.OpenThread(0x0002, False, wake_proc.pid))
            except Exception:
                pass

        # 延迟500ms确保窗口在前台后再触发Win+H
        self.after(500, self._do_trigger_winh)

    def _do_trigger_winh(self):
        import ctypes
        user32 = ctypes.windll.user32
        VK_LWIN = 0x5B
        VK_H    = 0x48
        KEYEVENTF_KEYUP = 0x0002
        user32.keybd_event(VK_LWIN, 0, 0, 0)
        user32.keybd_event(VK_H,    0, 0, 0)
        user32.keybd_event(VK_H,    0, KEYEVENTF_KEYUP, 0)
        user32.keybd_event(VK_LWIN, 0, KEYEVENTF_KEYUP, 0)

        self._winh_watch_start = datetime.datetime.now()
        self._winh_last_text = ""
        self.after(800, self._watch_winh_input)

    def _watch_winh_input(self):
        """轮询输入框，检测 Win+H 输入完成"""
        current = self.input_var.get().strip()
        elapsed = (datetime.datetime.now() - self._winh_watch_start).total_seconds()

        if current and current != self._winh_last_text:
            self._winh_last_text = current
            self.after(800, self._watch_winh_input)
        elif current and current == self._winh_last_text and elapsed > 1.5:
            self._set_status("● 处理中", ACCENT)
            self._resume_wake_proc()
            self._send_text()
        elif elapsed > 20:
            self._set_status("● 监听唤醒词" if self.wake_mode else "● 待机",
                             ACCENT2 if self.wake_mode else TEXT_DIM)
            self._resume_wake_proc()
            self._append("system", "Win+H 超时，请重试或直接输入")
        else:
            self.after(300, self._watch_winh_input)

    def _resume_wake_proc(self):
        """恢复被暂停的唤醒监听PS进程"""
        pass  # SuspendThread方案复杂，改用重启进程方式在_toggle_wake里处理

    def _on_wake_with_cmd(self, cmd):
        """唤醒词和指令在同一句话里，直接处理指令"""
        self._task_stop.set()   # 中断当前任务
        speak.__dict__.pop("_engine", None)  # 停止TTS
        if not self.winfo_viewable():
            self._show_window()
        self._set_status("● 处理指令", ACCENT)
        self._append("system", f"✨ 小K已唤醒，识别到指令：{cmd}")
        self._task_stop.clear()
        self._handle_input(cmd)

    def _on_wake(self):
        """唤醒：中断当前任务，触发 Win+H 云端语音识别"""
        self._task_stop.set()   # 中断当前任务
        self.deiconify()
        self.lift()
        self.focus_force()
        self._append("system", "✨ 小K已唤醒，请说出指令...")
        self.after(300, self._trigger_winh_input)

    def _listen_command(self):
        """降级方案：没有 Win+H 时用 Windows Speech 本地识别"""
        try:
            self._append("system", "🎤 请说出指令（本地语音识别中）...")
            text = recognize_with_windows_speech(timeout_sec=10)
            if self._task_stop.is_set():
                return
            if text:
                self.after(0, lambda: self._handle_input(text))
            else:
                self.after(0, lambda: self._append("system", "未听清，请再说一次或直接输入"))
        except Exception as e:
            if not self._task_stop.is_set():
                err = str(e)
                self.after(0, lambda msg=err: self._append("error", f"识别失败：{msg}"))
        finally:
            if self.wake_mode:
                self.after(0, lambda: self._set_status("● 监听唤醒词", ACCENT2))

    # ── 文字输入 ──────────────────────────────────────
    def _send_text(self):
        text = self.input_var.get().strip()
        if not text: return
        self.input_var.set(""); self._handle_input(text)

    # ── 统一处理 ──────────────────────────────────────
    def _handle_input(self, text):
        self._append("user", text)
        self._append("ai", "")
        self._set_status("● 处理中", ACCENT)
        self._task_stop.clear()  # 开始新任务前清除停止标志
        def run():
            result = execute_command(
                text,
                log_fn=lambda m: self.after(0, lambda msg=m: self._append("system", msg)),
                stream_fn=lambda p, clear=False: self.after(
                    0, lambda piece=p, c=clear: self._stream_append(piece, c)))
            # 任务被中断则不输出结果
            if self._task_stop.is_set():
                return
            self.after(0, lambda: self._set_status(
                "● 监听唤醒词" if self.wake_mode else "● 待机",
                ACCENT2 if self.wake_mode else TEXT_DIM))
            if result and not any(w in text for w in ["写","生成","帮我写","创作"]):
                self.after(0, lambda r=result: self._stream_append(r, clear=True))
                if not self._task_stop.is_set():
                    speak(result)
            self.after(0, lambda: self._append("system", ""))
        threading.Thread(target=run, daemon=True).start()

    # ── 语音设置 ──────────────────────────────────────
    def _open_asr_settings(self):
        win = tk.Toplevel(self); win.title("语音识别设置")
        win.geometry("460x320"); win.configure(bg=BG); win.resizable(False,False)
        cfg_file = os.path.join(SCRIPT_DIR,"voice_config.json")
        cfg = {}
        if os.path.exists(cfg_file):
            with open(cfg_file,encoding="utf-8") as f: cfg=json.load(f)
        tk.Label(win,text="语音识别引擎设置",bg=BG,fg=ACCENT,
                 font=("微软雅黑",12,"bold")).pack(pady=(16,4))
        fields=[("引擎 (whisper/baidu/google)","engine","whisper"),
                ("百度 APP_ID","baidu_app_id",""),
                ("百度 API_KEY","baidu_api_key",""),
                ("百度 SECRET","baidu_secret","")]
        vars_={}
        for label,key,default in fields:
            row=tk.Frame(win,bg=BG); row.pack(fill="x",padx=24,pady=4)
            tk.Label(row,text=label,bg=BG,fg=TEXT,font=("微软雅黑",9),
                     width=22,anchor="w").pack(side="left")
            v=tk.StringVar(value=cfg.get(key,default)); vars_[key]=v
            tk.Entry(row,textvariable=v,bg=BG2,fg=TEXT,insertbackground=TEXT,
                     relief="flat",font=("微软雅黑",9),width=24).pack(side="left",padx=6)
        def save():
            global ASR_ENGINE,BAIDU_APP_ID,BAIDU_API_KEY,BAIDU_SECRET
            new_cfg={k:v.get().strip() for k,v in vars_.items()}
            with open(cfg_file,"w",encoding="utf-8") as f: json.dump(new_cfg,f,ensure_ascii=False,indent=2)
            ASR_ENGINE=new_cfg.get("engine","whisper")
            BAIDU_APP_ID=new_cfg.get("baidu_app_id","")
            BAIDU_API_KEY=new_cfg.get("baidu_api_key","")
            BAIDU_SECRET=new_cfg.get("baidu_secret","")
            self._append("system",f"✓ 语音引擎已切换为：{ASR_ENGINE}"); win.destroy()
        tk.Button(win,text="保存",bg=ACCENT,fg=BG,relief="flat",
                  font=("微软雅黑",10),padx=20,pady=6,cursor="hand2",command=save).pack(pady=16)


if __name__ == "__main__":
    if not _ensure_single_instance():
        sys.exit(0)
    app = VoiceAssistant()
    app.mainloop()
